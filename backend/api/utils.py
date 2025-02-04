import os
from docx import Document

def replace_placeholders_in_docx(input_file, placeholders, username):
    """
    Replace placeholders in a Word document based on a dictionary of placeholders and save the new document
    in the user's subdirectory under /generated-documents.

    :param input_file: Path to the input Word document (.docx)
    :param placeholders: Dictionary of placeholder keys and replacement values
    :param username: Authenticated user's username or ID
    """
    # Define the output folder structure
    output_folder = os.path.join('generated-documents', username)

    # Ensure the output folder exists
    os.makedirs(output_folder, exist_ok=True)

    # Load the Word document
    doc = Document(input_file)

    # Iterate over all paragraphs in the document
    for paragraph in doc.paragraphs:
        full_text = "".join(run.text for run in paragraph.runs)
        updated_text = full_text
        for placeholder, value in placeholders.items():
            updated_text = updated_text.replace(f"<<{placeholder}>>", value)

        if full_text != updated_text:
            for run in paragraph.runs:
                run.text = ""  # Clear existing runs
            if paragraph.runs:
                paragraph.runs[0].text = updated_text
            else:
                paragraph.add_run(updated_text)

    # Iterate over all tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = "".join(run.text for run in paragraph.runs)
                    updated_text = full_text
                    for placeholder, value in placeholders.items():
                        updated_text = updated_text.replace(f"<<{placeholder}>>", value)

                    if full_text != updated_text:
                        for run in paragraph.runs:
                            run.text = ""  # Clear existing runs
                        if paragraph.runs:
                            paragraph.runs[0].text = updated_text
                        else:
                            paragraph.add_run(updated_text)

    # Define the output file path
    output_file = os.path.join(output_folder, os.path.basename(input_file))

    # Save the updated document
    doc.save(output_file)

    print(f"Document saved to: {output_file}")

    return output_file