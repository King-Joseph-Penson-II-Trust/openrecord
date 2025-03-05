from django.shortcuts import render, get_object_or_404
from django.contrib.auth.models import User
from rest_framework import generics, viewsets, permissions, status
from .serializers import UserSerializer, NoteSerializer, RecordSerializer, DocumentTemplateSerializer, GenerateDocumentsSerializer
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.views import APIView
from rest_framework.response import Response
from .models import Note, Record, DocumentTemplate
import logging
from django.utils import timezone
import requests
from api.utils import replace_placeholders_in_docx
from docx import Document
import json
from django.http import HttpResponse
from io import BytesIO
import os
from django.conf import settings

logger = logging.getLogger(__name__)


class NoteListCreate(generics.ListCreateAPIView):
    serializer_class = NoteSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        return Note.objects.filter(author=user)
    
    def perform_create(self, serializer):
        if serializer.is_valid():
            serializer.save(author=self.request.user)
        else:
            print(serializer.errors)

class NoteDelete(generics.DestroyAPIView):
    serializer_class = NoteSerializer
    permission_classes = [IsAuthenticated] #AllowAny

    def get_queryset(self):
        user = self.request.user
        return Note.objects.filter(author=user)

class CreateUserView(generics.CreateAPIView):
    queryset = User.objects.all()
    serializer_class = UserSerializer
    permission_classes = [AllowAny]

## Document Templates

class DocumentTemplateListView(generics.ListAPIView):
    queryset = DocumentTemplate.objects.all()
    serializer_class = DocumentTemplateSerializer
    permission_classes = [AllowAny]
    
class ScanPlaceholdersView(APIView):
    permission_classes = [AllowAny]

    def post(self, request, *args, **kwargs):
        file = request.FILES['file']
        doc = Document(file)
        placeholders = []

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                text = run.text
                if '<<' in text and '>>' in text:
                    placeholders.extend([f"<<{ph.split('>>')[0]}>>" for ph in text.split('<<')[1:]])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            text = run.text
                            if '<<' in text and '>>' in text:
                                placeholders.extend([f"<<{ph.split('>>')[0]}>>" for ph in text.split('<<')[1:]])

        # Remove duplicates while preserving order
        placeholders = list(dict.fromkeys(placeholders))

        return Response({'placeholders': placeholders}, status=status.HTTP_200_OK)




class UploadTemplateView(APIView):
    permission_classes = [AllowAny]

    def post(self, request, *args, **kwargs):
        file = request.FILES['file']
        name = request.data.get('name')
        description = request.data.get('description')
        placeholders = json.loads(request.data.get('placeholders'))

        # Define the file path
        templates_dir = 'templates'
        file_path = os.path.join(templates_dir, file.name)

        # Ensure the templates directory exists
        os.makedirs(templates_dir, exist_ok=True)

        # Save the file to the defined path
        with open(file_path, 'wb+') as destination:
            for chunk in file.chunks():
                destination.write(chunk)

        # Create the DocumentTemplate instance
        template = DocumentTemplate.objects.create(
            name=name,
            description=description,
            file=file_path,
            placeholders=placeholders
        )

        return Response({'message': 'Template uploaded successfully', 'template_id': template.id, 'template_path': template.file_path}, status=status.HTTP_201_CREATED)


def replace_placeholders_in_docx(input_file, placeholders, output_folder):
    """
    Replace placeholders in a Word document based on a JSON of placeholders and save the new document.

    :param input_file: Path to the input Word document (.docx)
    :param placeholders: Dictionary of placeholder keys and replacement values
    :param output_folder: Path to the folder where the updated document will be saved
    """
    # Ensure the output folder exists
    os.makedirs(output_folder, exist_ok=True)

    # Load the Word document
    doc = Document(input_file)

    # Iterate over all paragraphs in the document
    for paragraph in doc.paragraphs:
        # Combine all runs into a single string for placeholder replacement
        full_text = "".join(run.text for run in paragraph.runs)
        updated_text = full_text
        for placeholder, value in placeholders.items():
            updated_text = updated_text.replace(f"<<{placeholder}>>", value)

        # Update runs if the text was modified
        if full_text != updated_text:
            for run in paragraph.runs:
                run.text = ""  # Clear existing runs
            paragraph.runs[0].text = updated_text  # Set updated text in the first run

    # Iterate over all tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Combine all runs into a single string for placeholder replacement
                    full_text = "".join(run.text for run in paragraph.runs)
                    updated_text = full_text
                    for placeholder, value in placeholders.items():
                        updated_text = updated_text.replace(f"<<{placeholder}>>", value)

                    # Update runs if the text was modified
                    if full_text != updated_text:
                        for run in paragraph.runs:
                            run.text = ""  # Clear existing runs
                        paragraph.runs[0].text = updated_text  # Set updated text in the first run

    # Define the output file path
    output_file = os.path.join(output_folder, os.path.basename(input_file))

    # Save the updated document
    doc.save(output_file)

    print(f"Document saved to: {output_file}")
    
    return output_file

class ReplacePlaceholdersView(APIView):
    permission_classes = [AllowAny]

    def post(self, request, *args, **kwargs):
        template_id = request.data.get('template_id')
        replacements = json.loads(request.data.get('replacements'))

        # Retrieve the document template based on the provided template ID
        try:
            template = DocumentTemplate.objects.get(id=template_id)
        except DocumentTemplate.DoesNotExist:
            return Response({'error': 'Template not found'}, status=status.HTTP_404_NOT_FOUND)

        # Update the input file path to the correct directory
        input_file_path = os.path.join('templates', os.path.basename(template.file_path))


        # Define the output folder
        output_folder = os.path.join('generated_docs')

        # Replace placeholders in the document
        output_file_path = replace_placeholders_in_docx(input_file_path, replacements, output_folder)

        return Response({'message': 'Document generated successfully', 'output_file': output_file_path}, status=status.HTTP_200_OK)

class DeleteTemplateView(APIView):
    permission_classes = [AllowAny]

    def delete(self, request, *args, **kwargs):
        template_id = kwargs.get('pk')

        # Retrieve the document template based on the provided template ID
        try:
            template = DocumentTemplate.objects.get(id=template_id)
        except DocumentTemplate.DoesNotExist:
            return Response({'error': 'Template not found'}, status=status.HTTP_404_NOT_FOUND)

        # Delete the associated file from the filesystem
        if template.file and os.path.isfile(template.file.path):
            os.remove(template.file.path)

        # Delete the template from the database
        template.delete()

        return Response({'message': 'Template deleted successfully'}, status=status.HTTP_204_NO_CONTENT)

## Records

class RecordListCreate(generics.ListCreateAPIView):
    queryset = Record.objects.all()
    serializer_class = RecordSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        return Record.objects.filter(created_by=user)
    
    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

class RecordDetailUpdateDelete(generics.RetrieveUpdateDestroyAPIView):
    queryset = Record.objects.all()
    serializer_class = RecordSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        return Record.objects.filter(created_by=user)
    
    def perform_update(self, serializer):
        serializer.save()

class RecordListView(generics.ListAPIView):
    queryset = Record.objects.all()
    serializer_class = RecordSerializer
    permission_classes = [IsAuthenticated]

class RecordView(generics.RetrieveAPIView):
    queryset = Record.objects.all()
    serializer_class = RecordSerializer
    permission_classes = [AllowAny]

class RecordSearchView(generics.ListAPIView):
    serializer_class = RecordSerializer
    permission_classes = [AllowAny]

    def get_queryset(self):
        tracking_number = self.request.query_params.get('tracking_number', None)
        return_receipt = self.request.query_params.get('return_receipt', None)
        queryset = Record.objects.all()

        if tracking_number:
            queryset = queryset.filter(tracking_number=tracking_number)
        if return_receipt:
            queryset = queryset.filter(return_receipt=return_receipt)

        return queryset

    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        response = super().list(request, *args, **kwargs)

        for record in queryset:
            access_log = {
                'ip_address': request.META.get('REMOTE_ADDR'),
                'timestamp': timezone.now().isoformat(),
                'user_agent': request.META.get('HTTP_USER_AGENT'),
                'location': self.get_location(request.META.get('REMOTE_ADDR')),
            }
            record.access_logs.append(access_log)
            record.save()

        return response

    def get_location(self, ip_address):
        try:
            response = requests.get(f'https://ipinfo.io/{ip_address}/json')
            return response.json()
        except requests.RequestException:
            return {}